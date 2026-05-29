from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


FONT_TITLE = "方正小标宋_GBK"
FONT_BODY = "方正仿宋_GBK"
FONT_HEITI = "方正黑体_GBK"
FONT_KAITI = "方正楷体_GBK"
FONT_SONGTI = "宋体"
FONT_LATIN = "Times New Roman"
BODY_SIZE_PT = 16
FIRST_LINE_INDENT_PT = BODY_SIZE_PT * 2


def set_east_asia_font(run, font_name: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = FONT_LATIN if font_name == FONT_BODY else font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_LATIN if font_name == FONT_BODY else font_name)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN if font_name == FONT_BODY else font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    spacing = r_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        r_pr.append(spacing)
    spacing.set(qn("w:val"), "-4")


def set_para(paragraph, line_pt: float = 28, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_line: bool = False) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(line_pt)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(FIRST_LINE_INDENT_PT) if first_line else None


def add_text_para(doc, text: str, font: str = FONT_BODY, size: float = BODY_SIZE_PT,
                  bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  line_pt: float = 28, first_line: bool = False):
    p = doc.add_paragraph()
    set_para(p, line_pt=line_pt, alignment=align, first_line=first_line)
    run = p.add_run(text)
    set_east_asia_font(run, font, size, bold=bold)
    return p


def add_level2_para(doc, text: str):
    """Format `（一）标题。正文` with heading and body fonts in one paragraph."""
    p = doc.add_paragraph()
    set_para(p, first_line=True)
    match = re.match(r"^(（[一二三四五六七八九十]+）[^。；;:：]*。)(.*)$", text)
    if not match:
        run = p.add_run(text)
        set_east_asia_font(run, FONT_KAITI, BODY_SIZE_PT)
        return p

    heading, body = match.groups()
    head_run = p.add_run(heading)
    set_east_asia_font(head_run, FONT_KAITI, BODY_SIZE_PT)
    if body:
        body_run = p.add_run(body)
        set_east_asia_font(body_run, FONT_BODY, BODY_SIZE_PT)
    return p


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_page_setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(3.8)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)
    section.different_first_page_header_footer = False
    section.odd_and_even_pages_header_footer = True

    settings = doc.settings._element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def add_page_number(paragraph, align) -> None:
    set_para(paragraph, line_pt=14, alignment=align)
    run = paragraph.add_run("— ")
    set_east_asia_font(run, FONT_SONGTI, 14)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)
    end_run = paragraph.add_run(" —")
    set_east_asia_font(end_run, FONT_SONGTI, 14)


def setup_page_numbers(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    even_footer = section.even_page_footer
    if footer.paragraphs:
        p = footer.paragraphs[0]
        p.clear()
    else:
        p = footer.add_paragraph()
    add_page_number(p, WD_ALIGN_PARAGRAPH.RIGHT)
    if even_footer.paragraphs:
        ep = even_footer.paragraphs[0]
        ep.clear()
    else:
        ep = even_footer.add_paragraph()
    add_page_number(ep, WD_ALIGN_PARAGRAPH.LEFT)


def is_level1(text: str) -> bool:
    return bool(re.match(r"^[一二三四五六七八九十]+、", text))


def is_level2(text: str) -> bool:
    return bool(re.match(r"^（[一二三四五六七八九十]+）", text))


def is_level3(text: str) -> bool:
    return bool(re.match(r"^\d+\.\s*", text))


def is_level4(text: str) -> bool:
    return bool(re.match(r"^（\d+）", text))


def extract_formal_text(source: Path) -> list[str]:
    doc = Document(str(source))
    texts = [p.text.strip() for p in doc.paragraphs]
    texts = [t.replace("**", "") for t in texts if t.strip() and t.strip() != "---"]
    title_idx = 0
    for i, text in enumerate(texts):
        if re.search(r"(通知|报告|请示|函|通报|批复)$", text) and len(text) < 80:
            title_idx = i
            break
    return texts[title_idx:]


def build(source: Path, output: Path) -> None:
    texts = extract_formal_text(source)
    doc = Document(str(source))
    clear_body(doc)
    set_page_setup(doc)
    setup_page_numbers(doc)

    add_text_para(doc, texts[0], FONT_TITLE, 22, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=30)

    attachment_start = None
    for i, text in enumerate(texts):
        if i > 0 and text == "附件":
            attachment_start = i
            break
    body_texts = texts[1:attachment_start] if attachment_start else texts[1:]
    attachment_texts = texts[attachment_start:] if attachment_start else []

    for text in body_texts:
        if text.startswith("附件："):
            add_text_para(doc, text, FONT_BODY, BODY_SIZE_PT, align=WD_ALIGN_PARAGRAPH.LEFT)
        elif is_level1(text):
            add_text_para(doc, text, FONT_HEITI, BODY_SIZE_PT, first_line=True)
        elif is_level2(text):
            add_level2_para(doc, text)
        elif is_level3(text):
            fixed = re.sub(r"^(\d+\.)\s*", r"\1 ", text)
            add_text_para(doc, fixed, FONT_BODY, BODY_SIZE_PT, bold=True, first_line=True)
        elif is_level4(text):
            add_text_para(doc, text, FONT_BODY, BODY_SIZE_PT, first_line=True)
        elif re.match(r"^202\d年\d+月\d+日$", text):
            add_text_para(doc, text, FONT_BODY, BODY_SIZE_PT, align=WD_ALIGN_PARAGRAPH.RIGHT)
        elif re.match(r".*(办公室|委员会|局|处|科|部门)$", text) and len(text) <= 30:
            for _ in range(3):
                add_text_para(doc, "", FONT_BODY, BODY_SIZE_PT)
            add_text_para(doc, text, FONT_BODY, BODY_SIZE_PT, align=WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            add_text_para(doc, text, FONT_BODY, BODY_SIZE_PT, first_line=not text.endswith("："))

    if attachment_texts:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        set_page_setup(doc)
        add_text_para(doc, attachment_texts[0], FONT_HEITI, BODY_SIZE_PT, align=WD_ALIGN_PARAGRAPH.LEFT)
        add_text_para(doc, "", FONT_BODY, BODY_SIZE_PT)
        if len(attachment_texts) > 1:
            add_text_para(doc, attachment_texts[1], FONT_TITLE, 22, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=30)
        for text in attachment_texts[2:]:
            if is_level1(text):
                add_text_para(doc, text, FONT_HEITI, BODY_SIZE_PT, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True)
            else:
                add_text_para(doc, text, FONT_BODY, BODY_SIZE_PT, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.save(str(output))


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: format_gongwen_docx.py input.docx output.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
